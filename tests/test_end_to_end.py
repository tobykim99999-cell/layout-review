from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm

from layout_review_agent.agents import LayoutReviewCoordinator
from layout_review_agent.docx_format import get_paragraph_format
from layout_review_agent.llm import LLMConfig, OpenAICompatibleLLMClient, load_llm_config
from layout_review_agent.sample import create_bad_sample
from layout_review_agent.spec_normalizer import extract_json_object, normalize_spec_to_profile
from layout_review_agent.web import parse_multipart, render_profile_upload_page, render_upload_page, safe_filename, safe_profile_id


class FakeLLMClient:
    def __init__(self, response: str) -> None:
        self.response = response
        self.prompt = ""

    def complete(self, prompt: str) -> str:
        self.prompt = prompt
        return self.response


class EndToEndTest(unittest.TestCase):
    def test_web_helpers_parse_upload_form(self) -> None:
        boundary = "----layout-review-test"
        body = (
            f"--{boundary}\r\n"
            'Content-Disposition: form-data; name="profile"\r\n\r\n'
            "default_undergraduate\r\n"
            f"--{boundary}\r\n"
            'Content-Disposition: form-data; name="document"; filename="论文.docx"\r\n'
            "Content-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document\r\n\r\n"
        ).encode("utf-8") + b"docx-bytes" + f"\r\n--{boundary}--\r\n".encode("utf-8")

        fields, files = parse_multipart(f"multipart/form-data; boundary={boundary}", body)

        self.assertEqual(fields["profile"], "default_undergraduate")
        self.assertEqual(files["document"][0], "论文.docx")
        self.assertEqual(files["document"][1], b"docx-bytes")
        self.assertTrue(render_upload_page().startswith(b"<!doctype html>"))
        self.assertIn(b'name="llm_advice" checked', render_upload_page())
        self.assertIn("智能体正在执行任务".encode("utf-8"), render_upload_page())
        self.assertIn(b'class=\"js-work-form\"', render_upload_page())
        self.assertIn("上传学校论文格式规范".encode("utf-8"), render_profile_upload_page())
        self.assertEqual(safe_filename("../bad paper.docx"), "bad_paper.docx")
        self.assertEqual(safe_profile_id("XX 大学 本科 2026"), "xx_2026")

    def test_upload_page_hides_old_draft_profiles(self) -> None:
        page = render_upload_page(
            [
                {
                    "profile_id": "school_draft",
                    "display_name": "学校规范草稿",
                    "version": "2026.1",
                    "valid": True,
                    "errors": [],
                    "is_demo": False,
                    "is_template": True,
                    "is_draft": True,
                    "source_details": {
                        "document_name": "学校规范.docx",
                        "uploaded_file": "rule_profiles/_sources/学校规范.docx",
                    },
                    "rule_count": 0,
                }
            ]
        )

        self.assertIn("已隐藏 1 个旧草稿规则库".encode("utf-8"), page)
        self.assertNotIn(b"/profiles/normalize", page)
        self.assertNotIn("学校规范草稿".encode("utf-8"), page)

    def test_audit_and_safe_fix_generate_reports(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            input_path = create_bad_sample(root / "samples")
            output_dir = root / "reports"

            result = LayoutReviewCoordinator().audit(input_path, output_dir, fix_safe=True)

            self.assertGreater(result["summary"]["total_issues"], 0)
            self.assertGreater(result["summary"]["auto_fixable_issues"], 0)
            self.assertTrue(Path(result["document"]["fixed_path"]).exists())
            self.assertTrue(Path(result["reports"]["json"]).exists())
            self.assertTrue(Path(result["reports"]["xlsx"]).exists())
            self.assertTrue(Path(result["reports"]["html"]).exists())
            self.assertTrue(Path(result["reports"]["annotated_docx"]).exists())
            self.assertGreater(len(Document(result["reports"]["annotated_docx"]).comments), 0)
            self.assertTrue((output_dir / "iteration_insights.json").exists())
            self.assertIn("iteration", result)
            self.assertIn("shared_context", result)
            self.assertTrue((output_dir / "post_fix_result.json").exists())
            self.assertLessEqual(
                result["post_fix_summary"]["total_issues"],
                result["summary"]["total_issues"],
            )

    def test_safe_fix_only_updates_clear_body_paragraphs(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            profile_path = root / "school_profile.json"
            profile_path.write_text(
                """{
  "profile_id": "school_profile",
  "display_name": "测试学校规范",
  "version": "2026.1",
  "rules": [
    {
      "id": "body-paragraph",
      "category": "正文格式",
      "description": "正文段落格式",
      "selector": {"element_type": "paragraph", "exclude_empty": true, "min_chars": 8},
      "expected": {"line_spacing": 1.25, "first_line_indent_cm": 0.74},
      "severity": "minor",
      "status": "auto_fixable",
      "confidence": 0.96,
      "auto_fix": true,
      "safe_paragraph_auto_fix": true,
      "safe_fix_fields": ["line_spacing", "first_line_indent_cm"],
      "tolerance": 0.08,
      "suggestion": "按规范统一正文格式。"
    }
  ],
  "required_sections": []
}""",
                encoding="utf-8",
            )
            input_path = root / "thesis.docx"
            document = Document()
            document.add_paragraph("摘要")
            abstract_body = document.add_paragraph("本文研究品牌营销复兴策略，摘要段落不应该被正文规则自动改动。")
            abstract_body.paragraph_format.line_spacing = 1.0
            abstract_body.paragraph_format.first_line_indent = Cm(0)
            document.add_paragraph("Key words: brand; marketing; strategy")
            heading = document.add_paragraph("第一章 绪论")
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            body = document.add_paragraph("这是正文中的明显格式错误段落，应该由安全修复器调整行距和首行缩进。")
            body.paragraph_format.line_spacing = 1.0
            body.paragraph_format.first_line_indent = Cm(0)
            document.add_paragraph("参考文献")
            reference = document.add_paragraph("[1] Smith J. Marketing strategy research.")
            reference.paragraph_format.line_spacing = 1.0
            reference.paragraph_format.first_line_indent = Cm(0)
            document.save(input_path)

            result = LayoutReviewCoordinator(profile=str(profile_path)).audit(
                input_path,
                root / "reports",
                fix_safe=True,
            )
            fixed = Document(result["document"]["fixed_path"])

            self.assertAlmostEqual(fixed.paragraphs[1].paragraph_format.first_line_indent.cm, 0, places=2)
            self.assertEqual(fixed.paragraphs[1].paragraph_format.line_spacing, 1.0)
            self.assertAlmostEqual(fixed.paragraphs[4].paragraph_format.first_line_indent.cm, 0.74, places=2)
            self.assertEqual(fixed.paragraphs[4].paragraph_format.line_spacing, 1.25)
            self.assertAlmostEqual(fixed.paragraphs[6].paragraph_format.first_line_indent.cm, 0, places=2)
            self.assertEqual(fixed.paragraphs[6].paragraph_format.line_spacing, 1.0)
            self.assertEqual(result["safe_fix"]["body_bounds"], {"start": 3, "end": 5})

    def test_body_paragraph_audit_skips_front_matter_tables_and_back_matter(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            profile_path = root / "school_profile.json"
            profile_path.write_text(
                """{
  "profile_id": "school_profile",
  "display_name": "测试学校规范",
  "version": "2026.1",
  "rules": [
    {
      "id": "body-paragraph",
      "category": "正文格式",
      "description": "正文段落格式",
      "selector": {"element_type": "paragraph", "exclude_empty": true, "min_chars": 8},
      "expected": {"line_spacing": 1.25, "first_line_indent_cm": 0.74},
      "severity": "minor",
      "status": "auto_fixable",
      "confidence": 0.96,
      "auto_fix": true,
      "safe_paragraph_auto_fix": true,
      "safe_fix_fields": ["line_spacing", "first_line_indent_cm"],
      "tolerance": 0.08,
      "suggestion": "按规范统一正文格式。"
    }
  ],
  "required_sections": []
}""",
                encoding="utf-8",
            )
            input_path = root / "thesis_with_front_matter.docx"
            document = Document()
            cover = document.add_paragraph("福州大学本科生毕业设计（论文）诚信承诺书")
            cover.paragraph_format.line_spacing = 1.0
            cover.paragraph_format.first_line_indent = Cm(0)

            table = document.add_table(rows=3, cols=2)
            table.cell(0, 0).text = "学生姓名"
            table.cell(0, 1).text = "王赛寒"
            table.cell(1, 0).text = "毕业设计（论文）题目"
            table.cell(1, 1).text = "中文：电动车全铝车架结构设计与优化"
            table.cell(2, 0).text = "外文"
            table.cell(2, 1).text = "Structural Design and Optimization of All-Cast Aluminum Frame"
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.line_spacing = 1.0
                        paragraph.paragraph_format.first_line_indent = Cm(0)

            commitment_heading = document.add_paragraph("学生承诺")
            commitment_heading.paragraph_format.line_spacing = 1.0
            commitment = document.add_paragraph(
                "我承诺在毕业设计（论文）活动中遵守学校有关规定，将学术规范作为基本要求。"
            )
            commitment.paragraph_format.line_spacing = 1.0
            commitment.paragraph_format.first_line_indent = Cm(0)
            abstract_body = document.add_paragraph("本文研究车架结构优化方法，摘要段落不应该被正文规则批注。")
            abstract_body.paragraph_format.line_spacing = 1.0
            abstract_body.paragraph_format.first_line_indent = Cm(0)
            document.add_paragraph("关键词：车架；结构优化；有限元分析")
            heading = document.add_paragraph("第一章 绪论")
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            body = document.add_paragraph("这是正文中的明显格式错误段落，应该只对这一段生成正文格式批注。")
            body.paragraph_format.line_spacing = 1.0
            body.paragraph_format.first_line_indent = Cm(0)
            document.add_paragraph("参考文献")
            reference = document.add_paragraph("[1] Smith J. Structural design research.")
            reference.paragraph_format.line_spacing = 1.0
            reference.paragraph_format.first_line_indent = Cm(0)
            document.save(input_path)

            result = LayoutReviewCoordinator(profile=str(profile_path)).audit(
                input_path,
                root / "reports",
                fix_safe=False,
            )

            body_issues = [issue for issue in result["issues"] if issue["rule_id"] == "body-paragraph"]
            self.assertEqual({issue["location"]["element_id"] for issue in body_issues}, {"p-6"})
            self.assertEqual(len(body_issues), 2)
            self.assertEqual(result["summary"]["total_issues"], 1)
            self.assertEqual(result["shared_context"]["metrics"]["audit_body_bounds"], {"start": 5, "end": 7})

    def test_cjk_font_prefers_east_asia_font_slot(self) -> None:
        document = Document()
        paragraph = document.add_paragraph()
        run = paragraph.add_run("中文正文使用宋体，西文字体可以不同。")
        run.font.name = "Times New Roman"
        run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")

        self.assertEqual(get_paragraph_format(paragraph)["font_name"], "宋体")

    def test_reference_format_audit_only_checks_reference_section_entries(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            profile_path = root / "school_profile.json"
            profile_path.write_text(
                """{
  "profile_id": "school_profile",
  "display_name": "测试学校规范",
  "version": "2026.1",
  "rules": [
    {
      "id": "references-format",
      "category": "参考文献",
      "description": "参考文献条目格式",
      "selector": {"element_type": "paragraph", "exclude_empty": true, "text_regex": "^\\\\s*\\\\[?\\\\d+\\\\]?\\\\s*.+"},
      "expected": {"font_name": "宋体"},
      "severity": "major",
      "status": "manual_guided",
      "confidence": 0.94,
      "auto_fix": false,
      "tolerance": 0.08,
      "suggestion": "按学校规范统一参考文献条目格式。"
    }
  ],
  "required_sections": []
}""",
                encoding="utf-8",
            )
            input_path = root / "reference_scope.docx"
            document = Document()
            document.add_paragraph("2023年10月30日")
            document.add_paragraph("第一章 绪论")
            document.add_paragraph("1990年到2003年是李宁品牌的创立发展期，这不是参考文献。")
            table = document.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "60-70后"
            document.add_paragraph("参考文献")
            reference = document.add_paragraph("[1] Smith J. Marketing strategy research.")
            reference.runs[0].font.name = "Arial"
            document.save(input_path)

            result = LayoutReviewCoordinator(profile=str(profile_path)).audit(
                input_path,
                root / "reports",
                fix_safe=False,
            )

            reference_issues = [issue for issue in result["issues"] if issue["rule_id"] == "references-format"]
            self.assertEqual({issue["location"]["element_id"] for issue in reference_issues}, {"p-4"})
            self.assertEqual(result["summary"]["total_issues"], 1)

    def test_shared_llm_is_advisory_only(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            input_path = create_bad_sample(root / "samples")
            result = LayoutReviewCoordinator().audit(input_path, root / "reports", llm_advice=True)

            self.assertEqual(result["shared_llm"]["mode"], "shared_llm_service")
            self.assertFalse(result["shared_llm"]["affects_score"])
            self.assertFalse(result["shared_llm"]["affects_fixes"])
            self.assertIn("prompt", result["shared_llm"])
            self.assertIn("shared_llm_prompt", result["shared_context"]["artifacts"])

    def test_llm_config_has_required_fields(self) -> None:
        config = LLMConfig(
            provider="openai-compatible",
            base_url="https://api.example.com/v1/chat/completions",
            api_key="test-key",
            model="test-model",
            temperature=0.1,
            max_tokens=512,
        )
        client = OpenAICompatibleLLMClient(config=config)

        self.assertTrue(config.enabled)
        self.assertEqual(client.config.provider, "openai-compatible")
        self.assertEqual(client.config.base_url, "https://api.example.com/v1/chat/completions")
        self.assertEqual(client.config.api_key, "test-key")
        self.assertEqual(client.config.model, "test-model")
        self.assertEqual(client.config.temperature, 0.1)
        self.assertEqual(client.config.max_tokens, 512)

    def test_llm_config_can_load_from_file(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            config_path = Path(temp_dir) / "llm_config.json"
            config_path.write_text(
                """{
  "provider": "test-provider",
  "base_url": "https://api.example.com/v1",
  "api_key": "test-key",
  "model": "test-model",
  "temperature": 0.3,
  "max_tokens": 777
}""",
                encoding="utf-8",
            )

            config = load_llm_config(config_path)

            self.assertTrue(config.enabled)
            self.assertEqual(config.provider, "test-provider")
            self.assertEqual(config.chat_completions_url, "https://api.example.com/v1/chat/completions")
            self.assertEqual(config.masked()["api_key"], "***")

    def test_extract_json_object_from_fenced_llm_response(self) -> None:
        parsed = extract_json_object(
            """模型返回：
```json
{"profile_id": "school_a", "display_name": "A", "version": "1", "rules": [], "required_sections": []}
```
"""
        )

        self.assertEqual(parsed["profile_id"], "school_a")

    def test_spec_normalizer_generates_active_selectable_profile(self) -> None:
        fake_response = """{
  "profile_id": "ignored",
  "display_name": "ignored",
  "version": "ignored",
  "is_template": true,
  "rules": [
    {
      "id": "body-format",
      "category": "正文格式",
      "description": "正文格式",
      "selector": {"element_type": "paragraph", "exclude_empty": true, "min_chars": 8},
      "expected": {"font_name": "宋体", "font_size_pt": 12, "line_spacing": 1.5},
      "severity": "minor",
      "status": "auto_fixable",
      "confidence": 0.96,
      "auto_fix": true,
      "tolerance": 0.08,
      "suggestion": "按规范统一正文格式。"
    }
  ],
  "required_sections": [
    {"id": "required-abstract", "label": "摘要", "text_regex": "^摘要$", "severity": "major"}
  ]
}"""
        client = FakeLLMClient(fake_response)

        result = normalize_spec_to_profile(
            spec_text="正文采用宋体小四，1.5倍行距。论文应包含摘要。",
            profile_id="school_a_undergraduate_2026",
            display_name="School A 本科规范",
            version="2026.1",
            source={"type": "uploaded_school_spec", "document_name": "spec.docx"},
            llm_client=client,
        )

        self.assertFalse(result.profile["is_template"])
        self.assertFalse(result.profile["is_draft"])
        self.assertEqual(result.profile["profile_id"], "school_a_undergraduate_2026")
        self.assertEqual(len(result.profile["rules"]), 1)
        self.assertIn("当前审核引擎支持", client.prompt)

    def test_spec_normalizer_falls_back_when_llm_returns_no_json(self) -> None:
        client = FakeLLMClient("")

        result = normalize_spec_to_profile(
            spec_text="论文正文采用宋体小四号，1.5倍行距，首行缩进2字符，两端对齐。论文应包含摘要、关键词、目录、参考文献。",
            profile_id="school_local_2026",
            display_name="本地抽取规范",
            version="2026.1",
            source={"type": "uploaded_school_spec", "document_name": "spec.docx"},
            llm_client=client,
        )

        self.assertFalse(result.profile["is_template"])
        self.assertFalse(result.profile["is_draft"])
        self.assertGreaterEqual(len(result.profile["rules"]), 1)
        self.assertIn("LLM 未返回可用 JSON", result.notes[0])
        self.assertEqual(result.profile["rules"][0]["expected"]["font_name"], "宋体")
        self.assertEqual(result.profile["rules"][0]["status"], "auto_fixable")
        self.assertTrue(result.profile["rules"][0]["auto_fix"])
        self.assertEqual(
            result.profile["rules"][0]["safe_fix_fields"],
            ["line_spacing", "first_line_indent_cm", "space_before_pt", "space_after_pt"],
        )
        self.assertIn("^Key\\s*words?\\s*[:：]", result.profile["rules"][0]["selector"]["text_regex_not"])

    def test_batch_creates_summary(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            sample_dir = root / "samples"
            create_bad_sample(sample_dir)
            output_dir = root / "batch"

            summary = LayoutReviewCoordinator().batch(sample_dir, output_dir, fix_safe=False)

            self.assertEqual(summary["document_count"], 1)
            self.assertIn("iteration", summary)
            self.assertTrue((output_dir / "batch_summary.json").exists())
            self.assertTrue((output_dir / "batch_iteration_insights.json").exists())
            self.assertTrue(Path(summary["documents"][0]["report"]).exists())

    def test_long_term_memory_file_is_appended(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            input_path = create_bad_sample(root / "samples")
            memory_path = root / "memory" / "review_memory.jsonl"

            LayoutReviewCoordinator(memory_path=memory_path).audit(input_path, root / "reports")

            self.assertTrue(memory_path.exists())
            self.assertGreater(len(memory_path.read_text(encoding="utf-8").splitlines()), 0)


if __name__ == "__main__":
    unittest.main()
