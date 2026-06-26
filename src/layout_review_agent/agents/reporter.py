from __future__ import annotations

import html
import json
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill

from layout_review_agent.agents.base import Agent
from layout_review_agent.models import AgentRunContext, Issue, utc_now


class ReportWriterAgent(Agent[dict[str, Path]]):
    def __init__(self) -> None:
        super().__init__(
            agent_id="report_writer",
            description="Write JSON, Excel, and HTML audit reports.",
        )

    def run(self, context: AgentRunContext, result: dict[str, Any]) -> dict[str, Path]:
        trace = context.start_trace(self.agent_id, "write_reports")
        context.output_dir.mkdir(parents=True, exist_ok=True)
        json_path = context.output_dir / "result.json"
        xlsx_path = context.output_dir / "issues.xlsx"
        html_path = context.output_dir / "audit_report.html"
        annotated_path = context.output_dir / "annotated.docx"

        paths = {"json": json_path, "xlsx": xlsx_path, "html": html_path, "annotated_docx": annotated_path}
        for name, path in paths.items():
            context.shared.record_artifact(f"report_{name}", str(path))
        context.shared.observe(self.agent_id, "Reports written.", **{key: str(value) for key, value in paths.items()})
        result["reports"] = {key: str(value) for key, value in paths.items()}
        result["shared_context"] = context.shared.to_dict()
        serializable = self._to_serializable(result)
        json_path.write_text(json.dumps(serializable, ensure_ascii=False, indent=2), encoding="utf-8")
        self._write_excel(xlsx_path, result.get("issues", []))
        self._write_html(html_path, serializable)
        self._write_annotated_docx(annotated_path, context.input_path, result.get("issues", []), context)

        trace.finish("ok", "Reports written.", **{key: str(value) for key, value in paths.items()})
        return paths

    def _to_serializable(self, value: Any) -> Any:
        if isinstance(value, Issue):
            return value.to_dict()
        if hasattr(value, "to_dict"):
            return value.to_dict()
        if isinstance(value, Path):
            return str(value)
        if isinstance(value, dict):
            return {key: self._to_serializable(item) for key, item in value.items()}
        if isinstance(value, list):
            return [self._to_serializable(item) for item in value]
        return value

    def _write_excel(self, path: Path, issues: list[Issue]) -> None:
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "问题清单"
        headers = [
            "问题ID",
            "规则ID",
            "严重等级",
            "类别",
            "处理状态",
            "置信度",
            "位置",
            "字段",
            "实际值",
            "期望值",
            "建议",
        ]
        sheet.append(headers)
        for cell in sheet[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="305496")
        for issue in issues:
            sheet.append(
                [
                    issue.issue_id,
                    issue.rule_id,
                    issue.severity,
                    issue.category,
                    issue.status,
                    issue.confidence,
                    issue.location.get("element_id"),
                    issue.field,
                    str(issue.actual),
                    str(issue.expected),
                    issue.suggestion,
                ]
            )
        widths = [18, 22, 12, 14, 16, 10, 16, 18, 18, 18, 40]
        for index, width in enumerate(widths, start=1):
            sheet.column_dimensions[chr(64 + index)].width = width
        workbook.save(path)

    def _write_annotated_docx(
        self,
        path: Path,
        input_path: Path,
        issues: list[Issue],
        context: AgentRunContext,
    ) -> None:
        document = Document(str(input_path))
        grouped = self._group_issues_for_comments(issues)
        comments_added = 0
        skipped = 0
        for element_id, element_issues in grouped.items():
            runs = self._comment_runs_for_element(document, element_id)
            if not runs:
                skipped += len(element_issues)
                continue
            document.add_comment(
                runs,
                text=self._comment_text(element_issues),
                author="排版审核智能体",
                initials="LR",
            )
            comments_added += 1
        document.save(str(path))
        context.shared.record_metric("annotated_comments_added", comments_added)
        context.shared.record_metric("annotated_issues_skipped", skipped)

    def _group_issues_for_comments(self, issues: list[Issue]) -> dict[str, list[Issue]]:
        grouped: dict[str, list[Issue]] = {}
        for issue in issues:
            element_id = str(issue.location.get("element_id") or "document")
            grouped.setdefault(element_id, []).append(issue)
        return grouped

    def _comment_runs_for_element(self, document: Any, element_id: str) -> list[Any]:
        paragraph = self._paragraph_for_element(document, element_id)
        if paragraph is None:
            paragraph = self._first_non_empty_paragraph(document)
        if paragraph is None:
            return []
        runs = [run for run in paragraph.runs if run.text]
        return runs or list(paragraph.runs)

    def _paragraph_for_element(self, document: Any, element_id: str) -> Any | None:
        if element_id.startswith("p-"):
            try:
                index = int(element_id.split("-", 1)[1])
            except ValueError:
                return None
            if 0 <= index < len(document.paragraphs):
                return document.paragraphs[index]
            return None
        if element_id.startswith("table-"):
            parts = element_id.split("-")
            if len(parts) != 5:
                return None
            try:
                table_index = int(parts[1])
                row_index = int(parts[2][1:])
                cell_index = int(parts[3][1:])
                paragraph_index = int(parts[4][1:])
                return document.tables[table_index].rows[row_index].cells[cell_index].paragraphs[paragraph_index]
            except (ValueError, IndexError):
                return None
        return None

    def _first_non_empty_paragraph(self, document: Any) -> Any | None:
        for paragraph in document.paragraphs:
            if paragraph.text.strip() and paragraph.runs:
                return paragraph
        return None

    def _comment_text(self, issues: list[Issue]) -> str:
        lines = ["排版审核提示"]
        for index, issue in enumerate(issues[:8], start=1):
            lines.extend(
                [
                    "",
                    f"{index}. [{issue.severity}/{issue.status}] {issue.category}",
                    f"规则：{issue.rule_id}",
                    f"字段：{issue.field or '-'}",
                    f"实际：{issue.actual}",
                    f"期望：{issue.expected}",
                    f"建议：{issue.suggestion}",
                ]
            )
        if len(issues) > 8:
            lines.append(f"\n该位置还有 {len(issues) - 8} 个问题，请查看 issues.xlsx 完整清单。")
        return "\n".join(lines)

    def _write_html(self, path: Path, result: dict[str, Any]) -> None:
        summary = result["summary"]
        profile = result["profile"]
        issues = result.get("issues", [])
        quality = result.get("quality_gate", {})
        shared_llm = result.get("shared_llm")
        safe_fix = result.get("safe_fix") or {}
        applied_count = len(safe_fix.get("applied", []))
        skipped_count = len(safe_fix.get("skipped", []))
        iteration = result.get("iteration", {})
        rows = []
        for issue in issues:
            location = issue.get("location", {})
            rows.append(
                "<tr>"
                f"<td>{html.escape(issue.get('severity', ''))}</td>"
                f"<td>{html.escape(issue.get('category', ''))}</td>"
                f"<td>{html.escape(issue.get('status', ''))}</td>"
                f"<td>{html.escape(issue.get('rule_id', ''))}</td>"
                f"<td>{html.escape(str(location.get('element_id', '')))}</td>"
                f"<td>{html.escape(str(issue.get('field', '')))}</td>"
                f"<td>{html.escape(str(issue.get('actual', '')))}</td>"
                f"<td>{html.escape(str(issue.get('expected', '')))}</td>"
                f"<td>{html.escape(issue.get('suggestion', ''))}</td>"
                "</tr>"
            )
        html_text = f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <title>排版审核报告</title>
  <style>
    body {{ font-family: Arial, "Microsoft YaHei", sans-serif; margin: 32px; color: #1f2937; }}
    h1 {{ margin-bottom: 8px; }}
    .muted {{ color: #64748b; }}
    .cards {{ display: grid; grid-template-columns: repeat(4, minmax(120px, 1fr)); gap: 12px; margin: 24px 0; }}
    .card {{ border: 1px solid #d8dee9; border-radius: 6px; padding: 14px; }}
    .value {{ font-size: 28px; font-weight: 700; }}
    table {{ width: 100%; border-collapse: collapse; font-size: 13px; }}
    th, td {{ border: 1px solid #d8dee9; padding: 8px; vertical-align: top; }}
    th {{ background: #f1f5f9; text-align: left; }}
  </style>
</head>
<body>
  <h1>排版审核报告</h1>
  <div class="muted">规则：{html.escape(profile["display_name"])} / {html.escape(profile["version"])} · 生成时间：{utc_now()}</div>
  <div class="cards">
    <div class="card"><div class="muted">合规得分</div><div class="value">{summary["score"]}</div></div>
    <div class="card"><div class="muted">问题总数</div><div class="value">{summary["total_issues"]}</div></div>
    <div class="card"><div class="muted">可自动修复</div><div class="value">{summary["auto_fixable_issues"]}</div></div>
    <div class="card"><div class="muted">需人工复核</div><div class="value">{summary["manual_required_issues"]}</div></div>
  </div>
  <p><strong>质量门禁：</strong>{html.escape(str(quality.get("status", "not_run")))}；{html.escape(str(quality.get("message", "")))}</p>
  <p><strong>自动修复：</strong>实际应用 {applied_count} 项，跳过 {skipped_count} 项。段落级格式默认不自动覆盖，避免破坏原文档版式。</p>
  <p><strong>共享 LLM：</strong>{html.escape(str((shared_llm or {}).get("status", "not_run")))}；这是所有智能体可共用的大模型能力层，不是独立审核智能体，不影响规则判定。</p>
  <p><strong>智能迭代：</strong>{html.escape(str(iteration.get("total_issues", "not_run")))} 个问题已进入迭代画像，详见 iteration_insights.json。</p>
  <h2>问题清单</h2>
  <table>
    <thead><tr><th>等级</th><th>类别</th><th>状态</th><th>规则</th><th>位置</th><th>字段</th><th>实际值</th><th>期望值</th><th>建议</th></tr></thead>
    <tbody>{''.join(rows)}</tbody>
  </table>
</body>
</html>
"""
        path.write_text(html_text, encoding="utf-8")
