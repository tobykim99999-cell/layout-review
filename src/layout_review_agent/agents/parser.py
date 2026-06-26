from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document

from layout_review_agent.agents.base import Agent
from layout_review_agent.docx_format import get_paragraph_format, length_to_cm
from layout_review_agent.models import AgentRunContext, DocumentElement, ParsedDocument


class DocumentParserAgent(Agent[ParsedDocument]):
    def __init__(self) -> None:
        super().__init__(
            agent_id="document_parser",
            description="Parse DOCX structure, section setup, paragraph text, tables, and direct formatting.",
        )

    def run(self, context: AgentRunContext, path: str | Path | None = None) -> ParsedDocument:
        trace = context.start_trace(self.agent_id, "parse_docx")
        target_path = Path(path) if path else context.input_path
        document = Document(str(target_path))

        sections = [
            DocumentElement(
                element_id=f"section-{index}",
                element_type="section",
                text=f"Section {index + 1}",
                location={"element_id": f"section-{index}", "section_index": index},
                format={
                    "page_width_cm": length_to_cm(section.page_width),
                    "page_height_cm": length_to_cm(section.page_height),
                    "top_margin_cm": length_to_cm(section.top_margin),
                    "bottom_margin_cm": length_to_cm(section.bottom_margin),
                    "left_margin_cm": length_to_cm(section.left_margin),
                    "right_margin_cm": length_to_cm(section.right_margin),
                },
            )
            for index, section in enumerate(document.sections)
        ]

        elements: list[DocumentElement] = []
        for index, paragraph in enumerate(document.paragraphs):
            elements.append(self._paragraph_to_element(paragraph, f"p-{index}", index, "body"))

        table_paragraph_count = 0
        for table_index, table in enumerate(document.tables):
            for row_index, row in enumerate(table.rows):
                for cell_index, cell in enumerate(row.cells):
                    for paragraph_index, paragraph in enumerate(cell.paragraphs):
                        element_id = f"table-{table_index}-r{row_index}-c{cell_index}-p{paragraph_index}"
                        elements.append(
                            self._paragraph_to_element(
                                paragraph,
                                element_id,
                                paragraph_index,
                                "table",
                                {
                                    "table_index": table_index,
                                    "row_index": row_index,
                                    "cell_index": cell_index,
                                },
                            )
                        )
                        table_paragraph_count += 1

        parsed = ParsedDocument(
            path=target_path,
            sections=sections,
            elements=elements,
            metadata={
                "paragraph_count": len(document.paragraphs),
                "table_count": len(document.tables),
                "table_paragraph_count": table_paragraph_count,
                "section_count": len(document.sections),
            },
        )
        context.shared.record_artifact("last_parsed_docx", str(target_path))
        context.shared.record_metric("parsed_elements", len(elements))
        context.shared.record_metric("parsed_sections", len(sections))
        context.shared.observe(
            self.agent_id,
            "DOCX parsed into shared review state.",
            path=str(target_path),
            elements=len(elements),
            sections=len(sections),
        )
        trace.finish(
            "ok",
            f"Parsed {len(elements)} elements from DOCX.",
            elements=len(elements),
            sections=len(sections),
        )
        return parsed

    def _paragraph_to_element(
        self,
        paragraph: Any,
        element_id: str,
        paragraph_index: int,
        scope: str,
        extra_location: dict[str, Any] | None = None,
    ) -> DocumentElement:
        location = {
            "element_id": element_id,
            "scope": scope,
            "paragraph_index": paragraph_index,
            "preview": " ".join(paragraph.text.split())[:80],
        }
        if extra_location:
            location.update(extra_location)
        return DocumentElement(
            element_id=element_id,
            element_type="paragraph",
            text=paragraph.text,
            location=location,
            style_name=paragraph.style.name if paragraph.style is not None else None,
            format=get_paragraph_format(paragraph),
        )
