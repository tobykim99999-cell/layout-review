from __future__ import annotations

import shutil
from pathlib import Path
from typing import Any

from docx import Document

from layout_review_agent.agents.base import Agent
from layout_review_agent.document_scope import (
    find_body_bounds_from_paragraphs,
    is_inside_body_bounds,
    looks_like_non_body_paragraph,
    paragraph_index_from_element_id,
)
from layout_review_agent.docx_format import apply_paragraph_field, apply_section_field
from layout_review_agent.models import AgentRunContext, Issue


class SafeFixerAgent(Agent[dict[str, Any]]):
    DEFAULT_SAFE_PARAGRAPH_FIELDS = {
        "line_spacing",
        "first_line_indent_cm",
        "space_before_pt",
        "space_after_pt",
    }

    def __init__(self, confidence_threshold: float = 0.94) -> None:
        super().__init__(
            agent_id="safe_fixer",
            description="Apply only high-confidence deterministic DOCX formatting fixes.",
        )
        self.confidence_threshold = confidence_threshold

    def run(self, context: AgentRunContext, issues: list[Issue], fixed_path: Path) -> dict[str, Any]:
        trace = context.start_trace(self.agent_id, "safe_fix")
        fixed_path.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(context.input_path, fixed_path)
        document = Document(str(fixed_path))
        body_bounds = find_body_bounds_from_paragraphs(list(document.paragraphs))

        applied: list[dict[str, Any]] = []
        skipped: list[dict[str, Any]] = []
        for issue in issues:
            if issue.status != "auto_fixable" or issue.confidence < self.confidence_threshold:
                skipped.append({"issue_id": issue.issue_id, "reason": "not_auto_fixable_or_low_confidence"})
                continue
            strategy = issue.fix_strategy
            target = strategy.get("target")
            element_id = strategy.get("element_id", "")
            field = strategy.get("field")
            value = strategy.get("value")
            ok = False
            if target == "paragraph" and field:
                if not strategy.get("allow_paragraph_fix", False):
                    record = {
                        "issue_id": issue.issue_id,
                        "rule_id": issue.rule_id,
                        "element_id": element_id,
                        "field": field,
                        "value": value,
                        "reason": "paragraph_auto_fix_requires_explicit_safe_paragraph_auto_fix",
                    }
                    skipped.append(record)
                    continue
                safe_fields = set(strategy.get("safe_fix_fields") or self.DEFAULT_SAFE_PARAGRAPH_FIELDS)
                if field not in safe_fields:
                    skipped.append(
                        {
                            "issue_id": issue.issue_id,
                            "rule_id": issue.rule_id,
                            "element_id": element_id,
                            "field": field,
                            "value": value,
                            "reason": "paragraph_field_not_in_safe_fix_fields",
                        }
                    )
                    continue
                if element_id.startswith("table-"):
                    skipped.append(
                        {
                            "issue_id": issue.issue_id,
                            "rule_id": issue.rule_id,
                            "element_id": element_id,
                            "field": field,
                            "value": value,
                            "reason": "table_paragraph_auto_fix_skipped",
                        }
                    )
                    continue
                paragraph = self._resolve_paragraph(document, element_id)
                if paragraph is not None:
                    if self._paragraph_is_safe_for_auto_fix(paragraph, issue.rule_id, element_id, body_bounds):
                        ok = apply_paragraph_field(paragraph, field, value)
                    else:
                        skipped.append(
                            {
                                "issue_id": issue.issue_id,
                                "rule_id": issue.rule_id,
                                "element_id": element_id,
                                "field": field,
                                "value": value,
                                "reason": "paragraph_shape_not_safe_for_auto_fix",
                            }
                        )
                        continue
            elif target == "section" and field:
                section = self._resolve_section(document, element_id)
                if section is not None:
                    ok = apply_section_field(section, field, value)

            record = {
                "issue_id": issue.issue_id,
                "rule_id": issue.rule_id,
                "element_id": element_id,
                "field": field,
                "value": value,
            }
            if ok:
                applied.append(record)
            else:
                record["reason"] = "target_not_found_or_field_not_supported"
                skipped.append(record)

        document.save(str(fixed_path))
        context.shared.record_artifact("fixed_docx", str(fixed_path))
        context.shared.record_metric("safe_fixes_applied", len(applied))
        context.shared.record_metric("safe_fixes_skipped", len(skipped))
        context.shared.decide(
            self.agent_id,
            "apply_safe_fixes",
            "Only auto_fixable issues above the confidence threshold were modified.",
            confidence_threshold=self.confidence_threshold,
            applied=len(applied),
            skipped=len(skipped),
        )
        trace.finish(
            "ok",
            f"Applied {len(applied)} safe fixes.",
            applied=len(applied),
            skipped=len(skipped),
            fixed_path=str(fixed_path),
        )
        return {
            "fixed_path": fixed_path,
            "applied": applied,
            "skipped": skipped,
            "body_bounds": body_bounds,
        }

    def _resolve_paragraph(self, document: Any, element_id: str) -> Any | None:
        if element_id.startswith("p-"):
            index = int(element_id.split("-", 1)[1])
            if 0 <= index < len(document.paragraphs):
                return document.paragraphs[index]
            return None
        if element_id.startswith("table-"):
            parts = element_id.split("-")
            if len(parts) != 5:
                return None
            table_index = int(parts[1])
            row_index = int(parts[2][1:])
            cell_index = int(parts[3][1:])
            paragraph_index = int(parts[4][1:])
            try:
                return document.tables[table_index].rows[row_index].cells[cell_index].paragraphs[paragraph_index]
            except IndexError:
                return None
        return None

    def _paragraph_is_safe_for_auto_fix(
        self,
        paragraph: Any,
        rule_id: str,
        element_id: str,
        body_bounds: dict[str, int | None],
    ) -> bool:
        text = paragraph.text.strip()
        if len(text) < 8:
            return False
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if rule_id == "body-paragraph":
            index = paragraph_index_from_element_id(element_id)
            if not is_inside_body_bounds(index, body_bounds):
                return False
            if looks_like_non_body_paragraph(text, style_name):
                return False
            if "Heading" in style_name or "标题" in style_name:
                return False
            if str(paragraph.alignment) in {"CENTER (1)", "RIGHT (2)"}:
                return False
        for run in paragraph.runs:
            if not run.text.strip() or run.font.size is None:
                continue
            if run.font.size.pt and run.font.size.pt > 18:
                return False
        return True

    def _resolve_section(self, document: Any, element_id: str) -> Any | None:
        if not element_id.startswith("section-"):
            return None
        index = int(element_id.split("-", 1)[1])
        if 0 <= index < len(document.sections):
            return document.sections[index]
        return None
