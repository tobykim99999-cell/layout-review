from __future__ import annotations

import io
import json
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document

from layout_review_agent.llm import LLMClient
from layout_review_agent.rules import validate_profile_data

SUPPORTED_EXPECTED_FIELDS = {
    "page_width_cm",
    "page_height_cm",
    "top_margin_cm",
    "bottom_margin_cm",
    "left_margin_cm",
    "right_margin_cm",
    "font_name",
    "font_size_pt",
    "bold",
    "italic",
    "alignment",
    "line_spacing",
    "first_line_indent_cm",
    "space_before_pt",
    "space_after_pt",
}

SUPPORTED_ELEMENT_TYPES = {"section", "paragraph"}
VALID_STATUSES = {"auto_fixable", "manual_guided", "manual_required"}
MAX_SPEC_CHARS = 22000
FONT_SIZE_NAME_TO_PT = {
    "一": 26,
    "小一": 24,
    "二": 22,
    "小二": 18,
    "三": 16,
    "小三": 15,
    "四": 14,
    "小四": 12,
    "五": 10.5,
    "小五": 9,
    "六": 7.5,
}
CHINESE_FONTS = ("宋体", "黑体", "楷体", "楷体_GB2312", "仿宋", "微软雅黑", "Times New Roman")


class SpecNormalizationError(ValueError):
    """Raised when an uploaded school spec cannot be turned into an active rule profile."""


@dataclass(frozen=True)
class SpecNormalizationResult:
    profile: dict[str, Any]
    prompt: str
    response: str
    source_text_chars: int
    used_text_chars: int
    notes: list[str]


def extract_spec_text(filename: str, payload: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".docx":
        return _extract_docx_text(payload)
    if suffix in {".txt", ".md"}:
        return _decode_text(payload)
    if suffix == ".pdf":
        return _extract_pdf_text(payload)
    raise SpecNormalizationError("仅支持从 .docx、.txt、.md、.pdf 规范中抽取文本。")


def normalize_spec_to_profile(
    *,
    spec_text: str,
    profile_id: str,
    display_name: str,
    version: str,
    source: dict[str, Any],
    llm_client: LLMClient | None,
) -> SpecNormalizationResult:
    cleaned_text = _compact_spec_text(spec_text)
    if not cleaned_text.strip():
        raise SpecNormalizationError("规范文件没有抽取到可用文字，无法自动生成规则库。")

    prompt = build_profile_generation_prompt(
        spec_text=cleaned_text,
        profile_id=profile_id,
        display_name=display_name,
        version=version,
        source=source,
    )
    response = ""
    if llm_client is None:
        return build_local_profile_result(
            spec_text=spec_text,
            cleaned_text=cleaned_text,
            profile_id=profile_id,
            display_name=display_name,
            version=version,
            source=source,
            prompt=prompt,
            response=response,
            fallback_reason="LLM 未配置，已使用本地确定性规则抽取。",
        )

    try:
        response = llm_client.complete(prompt)
    except Exception as exc:
        return build_local_profile_result(
            spec_text=spec_text,
            cleaned_text=cleaned_text,
            profile_id=profile_id,
            display_name=display_name,
            version=version,
            source=source,
            prompt=prompt,
            response=response,
            fallback_reason=f"LLM 调用失败，已使用本地确定性规则抽取：{exc}",
        )
    try:
        raw = extract_json_object(response)
        profile, notes = sanitize_generated_profile(
            raw,
            profile_id=profile_id,
            display_name=display_name,
            version=version,
            source=source,
        )
        errors = validate_profile_data(profile)
        if errors:
            raise SpecNormalizationError("生成的规则库未通过结构校验：" + "；".join(errors))
        if not profile.get("rules"):
            raise SpecNormalizationError("生成的规则库没有可执行 rules，不能作为可选审核规则。")
    except SpecNormalizationError as exc:
        return build_local_profile_result(
            spec_text=spec_text,
            cleaned_text=cleaned_text,
            profile_id=profile_id,
            display_name=display_name,
            version=version,
            source=source,
            prompt=prompt,
            response=response,
            fallback_reason=f"LLM 未返回可用 JSON，已使用本地确定性规则抽取：{exc}",
        )
    return SpecNormalizationResult(
        profile=profile,
        prompt=prompt,
        response=response,
        source_text_chars=len(spec_text),
        used_text_chars=len(cleaned_text),
        notes=notes,
    )


def build_local_profile_result(
    *,
    spec_text: str,
    cleaned_text: str,
    profile_id: str,
    display_name: str,
    version: str,
    source: dict[str, Any],
    prompt: str,
    response: str,
    fallback_reason: str,
) -> SpecNormalizationResult:
    profile, notes = build_local_rule_profile(
        cleaned_text,
        profile_id=profile_id,
        display_name=display_name,
        version=version,
        source=source,
    )
    notes.insert(0, fallback_reason)
    errors = validate_profile_data(profile)
    if errors:
        raise SpecNormalizationError("本地规则库生成结果未通过结构校验：" + "；".join(errors))
    if not profile.get("rules"):
        raise SpecNormalizationError(f"{fallback_reason}；但未从规范中抽取到当前引擎可执行的格式规则。")
    return SpecNormalizationResult(
        profile=profile,
        prompt=prompt,
        response=response,
        source_text_chars=len(spec_text),
        used_text_chars=len(cleaned_text),
        notes=notes,
    )


def build_profile_generation_prompt(
    *,
    spec_text: str,
    profile_id: str,
    display_name: str,
    version: str,
    source: dict[str, Any],
) -> str:
    source_json = json.dumps(source, ensure_ascii=False, indent=2)
    supported_fields = ", ".join(sorted(SUPPORTED_EXPECTED_FIELDS))
    return f"""你是毕业论文排版审核系统的“规则库规范化器”。你的任务是把学校官方论文格式规范转换成严格 JSON。

硬性要求：
1. 只输出一个 JSON 对象，不要 Markdown，不要解释文字。
2. 只能根据规范原文明确写出的要求生成规则；不确定、缺失、需要分页渲染判断的内容不要猜。
3. 只能生成当前审核引擎支持的 expected 字段：{supported_fields}。
4. selector.element_type 只能是 "section" 或 "paragraph"。
5. 自动修复只允许用于字体、字号、粗斜、对齐、行距、缩进、段前段后、页边距、纸张大小这类确定性格式项。
6. 目录页码、真实页码连续性、跨页表格、图片压缩质量、域代码、复杂页眉页脚等不能直接用当前引擎判断的内容，写入 manual_review_notes，不要放入 expected。
7. 字号换算请用：一号=26，小一=24，二号=22，小二=18，三号=16，小三=15，四号=14，小四=12，五号=10.5，小五=9，六号=7.5。
8. alignment 只能用 "left"、"center"、"right"、"justify"。
9. status 只能用 "auto_fixable"、"manual_guided"、"manual_required"。
10. regex 必须是 JSON 字符串可用的正则，反斜杠要正确转义。

目标规则库元数据：
profile_id: {profile_id}
display_name: {display_name}
version: {version}
source:
{source_json}

输出 JSON 结构必须类似：
{{
  "profile_id": "{profile_id}",
  "display_name": "{display_name}",
  "version": "{version}",
  "is_demo": false,
  "is_template": false,
  "is_draft": false,
  "source": {{...}},
  "description": "根据上传的学校官方规范自动生成，后续仍可人工复核。",
  "rules": [
    {{
      "id": "body-paragraph",
      "category": "正文格式",
      "description": "正文段落格式",
      "selector": {{
        "element_type": "paragraph",
        "exclude_empty": true,
        "min_chars": 8,
        "text_regex_not": ["^摘要$", "^Abstract$", "^关键词[:：]", "^Key\\\\s*words?\\\\s*[:：]", "^Keywords\\\\s*[:：]", "^目录$", "^参考文献$"]
      }},
      "expected": {{
        "font_name": "宋体",
        "font_size_pt": 12,
        "alignment": "justify",
        "line_spacing": 1.5,
        "first_line_indent_cm": 0.74
      }},
      "severity": "minor",
      "status": "auto_fixable",
      "confidence": 0.95,
      "auto_fix": true,
      "tolerance": 0.08,
      "suggestion": "按学校规范统一正文格式。"
    }}
  ],
  "required_sections": [
    {{
      "id": "required-abstract",
      "category": "特殊模块",
      "label": "摘要",
      "text_regex": "^摘要$",
      "severity": "major",
      "suggestion": "论文应包含摘要模块。"
    }}
  ],
  "manual_review_notes": ["无法用 DOCX 静态解析可靠判断的规范项。"]
}}

学校规范原文：
{spec_text}
"""


def extract_json_object(text: str) -> dict[str, Any]:
    candidates: list[str] = []
    stripped = text.strip()
    candidates.append(stripped)
    for match in re.finditer(r"```(?:json)?\s*(.*?)```", text, flags=re.IGNORECASE | re.DOTALL):
        candidates.append(match.group(1).strip())
    balanced = _find_balanced_json_object(text)
    if balanced:
        candidates.append(balanced)

    last_error: Exception | None = None
    for candidate in candidates:
        if not candidate:
            continue
        try:
            parsed = json.loads(candidate)
        except json.JSONDecodeError as exc:
            last_error = exc
            continue
        if not isinstance(parsed, dict):
            raise SpecNormalizationError("模型输出 JSON 顶层必须是对象。")
        return parsed
    message = f"无法从模型输出中解析 JSON：{last_error}" if last_error else "模型没有返回 JSON。"
    raise SpecNormalizationError(message)


def sanitize_generated_profile(
    raw: dict[str, Any],
    *,
    profile_id: str,
    display_name: str,
    version: str,
    source: dict[str, Any],
) -> tuple[dict[str, Any], list[str]]:
    notes: list[str] = []
    profile = dict(raw)
    profile["profile_id"] = profile_id
    profile["display_name"] = str(profile.get("display_name") or display_name)
    profile["version"] = str(profile.get("version") or version)
    profile["is_demo"] = False
    profile["is_template"] = False
    profile["is_draft"] = False
    profile["source"] = {**source, **dict(profile.get("source") or {})}
    profile.setdefault("description", "根据上传的学校官方规范自动生成。")

    cleaned_rules: list[dict[str, Any]] = []
    for index, rule in enumerate(profile.get("rules") or []):
        if not isinstance(rule, dict):
            notes.append(f"跳过 rules[{index}]：不是对象。")
            continue
        selector = dict(rule.get("selector") or {})
        element_type = selector.get("element_type")
        if element_type not in SUPPORTED_ELEMENT_TYPES:
            notes.append(f"跳过规则 {rule.get('id', index)}：element_type 不受支持。")
            continue

        expected = dict(rule.get("expected") or {})
        unsupported = sorted(set(expected) - SUPPORTED_EXPECTED_FIELDS)
        for field in unsupported:
            expected.pop(field, None)
        if unsupported:
            notes.append(f"规则 {rule.get('id', index)} 已移除不支持字段：{', '.join(unsupported)}。")
        if not expected:
            notes.append(f"跳过规则 {rule.get('id', index)}：没有当前引擎可执行的 expected 字段。")
            continue
        if not _selector_regex_valid(selector):
            notes.append(f"跳过规则 {rule.get('id', index)}：selector 正则无效。")
            continue

        rule_id = str(rule.get("id") or f"generated-rule-{index + 1}")
        status = rule.get("status")
        if status not in VALID_STATUSES:
            status = "manual_guided"
        clean_rule = {
            **rule,
            "id": _safe_rule_id(rule_id, f"generated-rule-{index + 1}"),
            "category": str(rule.get("category") or "通用格式"),
            "description": str(rule.get("description") or rule_id),
            "selector": selector,
            "expected": expected,
            "severity": str(rule.get("severity") or "minor"),
            "status": status,
            "confidence": _clamp_float(rule.get("confidence"), 0.9, 0.0, 1.0),
            "auto_fix": bool(rule.get("auto_fix", status == "auto_fixable")) and status == "auto_fixable",
            "tolerance": _clamp_float(rule.get("tolerance"), 0.08, 0.0, 2.0),
            "suggestion": str(rule.get("suggestion") or "请按学校规范修正。"),
        }
        if element_type == "paragraph":
            if clean_rule["auto_fix"] and _is_body_rule(clean_rule):
                clean_rule["safe_paragraph_auto_fix"] = True
                clean_rule["safe_fix_fields"] = _safe_paragraph_fix_fields(rule.get("safe_fix_fields"))
                clean_rule["selector"] = _with_default_body_exclusions(selector)
            elif clean_rule["auto_fix"]:
                clean_rule["status"] = "manual_guided"
                clean_rule["auto_fix"] = False
                notes.append(f"规则 {clean_rule['id']} 为非正文段落规则，已关闭自动修复以避免误改。")
        cleaned_rules.append(clean_rule)

    seen_rule_ids: set[str] = set()
    for rule in cleaned_rules:
        base_id = rule["id"]
        candidate = base_id
        suffix = 2
        while candidate in seen_rule_ids:
            candidate = f"{base_id}-{suffix}"
            suffix += 1
        rule["id"] = candidate
        seen_rule_ids.add(candidate)

    cleaned_sections: list[dict[str, Any]] = []
    for index, section in enumerate(profile.get("required_sections") or []):
        if not isinstance(section, dict):
            notes.append(f"跳过 required_sections[{index}]：不是对象。")
            continue
        pattern = str(section.get("text_regex") or "")
        if not pattern or not _regex_valid(pattern):
            notes.append(f"跳过必备模块 {section.get('id', index)}：text_regex 为空或无效。")
            continue
        section_id = _safe_rule_id(str(section.get("id") or f"required-section-{index + 1}"), f"required-section-{index + 1}")
        cleaned_sections.append(
            {
                **section,
                "id": section_id,
                "category": str(section.get("category") or "结构完整性"),
                "label": str(section.get("label") or section_id),
                "text_regex": pattern,
                "severity": str(section.get("severity") or "major"),
                "suggestion": str(section.get("suggestion") or "请补充必备论文模块。"),
            }
        )

    profile["rules"] = cleaned_rules
    profile["required_sections"] = cleaned_sections
    profile["normalization"] = {
        "method": "llm_generated_then_schema_sanitized",
        "supported_expected_fields": sorted(SUPPORTED_EXPECTED_FIELDS),
        "notes": notes,
    }
    return profile, notes


def build_local_rule_profile(
    spec_text: str,
    *,
    profile_id: str,
    display_name: str,
    version: str,
    source: dict[str, Any],
) -> tuple[dict[str, Any], list[str]]:
    notes: list[str] = ["本地抽取器只生成规范原文中能明确识别且当前引擎可执行的格式规则。"]
    rules: list[dict[str, Any]] = []

    section_expected = _extract_section_expected(spec_text)
    if section_expected:
        rules.append(
            _make_rule(
                rule_id="page-setup",
                category="页面设置",
                description="页面尺寸与页边距",
                selector={"element_type": "section"},
                expected=section_expected,
                severity="major",
                suggestion="按学校规范统一页面尺寸与页边距。",
                auto_fix=True,
            )
        )

    global_line_spacing = _extract_global_line_spacing(spec_text)
    body_context = _context_for_best_label(spec_text, ("论文正文", "正文内容", "正文文字", "正文段落", "正文格式", "正文"))
    body_expected = _extract_paragraph_expected(body_context)
    if body_expected and global_line_spacing is not None:
        body_expected.setdefault("line_spacing", global_line_spacing)
    if body_expected:
        body_expected.setdefault("alignment", "justify")
        rules.append(
            _make_rule(
                rule_id="body-paragraph",
                category="正文格式",
                description="正文段落格式",
                selector={
                    "element_type": "paragraph",
                    "exclude_empty": True,
                    "min_chars": 8,
                    "text_regex_not": [
                        "^(第[一二三四五六七八九十0-9]+章|[一二三四五六七八九十]+、)\\s*.+",
                        "^(\\d+\\.\\d+|（[一二三四五六七八九十]+）)\\s*.+",
                        "^摘要$",
                        "^Abstract$",
                        "^关键词[:：]",
                        "^Key\\s*words?\\s*[:：]",
                        "^Keywords\\s*[:：]",
                        "^目录$",
                        "^Contents$",
                        "^参考文献$",
                        "^致谢$",
                        "^(图|表)\\s*\\d+",
                    ],
                },
                expected=body_expected,
                severity="minor",
                suggestion="按学校规范统一正文格式。",
                auto_fix=True,
                safe_paragraph_auto_fix=True,
                safe_fix_fields=["line_spacing", "first_line_indent_cm", "space_before_pt", "space_after_pt"],
            )
        )

    heading_specs = [
        (
            "level-1-heading",
            "一级标题格式",
            ("一级标题", "章标题", "章名", "第一层次标题"),
            "^(第[一二三四五六七八九十0-9]+章|[一二三四五六七八九十]+、)\\s*.+",
        ),
        (
            "level-2-heading",
            "二级标题格式",
            ("二级标题", "节标题", "第二层次标题"),
            "^(\\d+\\.\\d+|（[一二三四五六七八九十]+）)\\s*.+",
        ),
        (
            "level-3-heading",
            "三级标题格式",
            ("三级标题", "第三层次标题"),
            "^(\\d+\\.\\d+\\.\\d+|[0-9]+[).、])\\s*.+",
        ),
    ]
    for rule_id, description, keywords, pattern in heading_specs:
        expected = _extract_paragraph_expected(_context_for_best_label(spec_text, keywords))
        if expected and global_line_spacing is not None:
            expected.setdefault("line_spacing", global_line_spacing)
        if not expected:
            continue
        rules.append(
            _make_rule(
                rule_id=rule_id,
                category="标题体系",
                description=description,
                selector={"element_type": "paragraph", "exclude_empty": True, "text_regex": pattern},
                expected=expected,
                severity="major",
                suggestion=f"按学校规范统一{description}。",
                auto_fix=False,
            )
        )

    caption_context = _context_for_best_label(spec_text, ("图题、表题", "图题", "表题", "题注", "图名", "表名"))
    caption_expected = _extract_paragraph_expected(caption_context)
    if caption_expected and global_line_spacing is not None:
        caption_expected.setdefault("line_spacing", global_line_spacing)
    if caption_expected:
        rules.append(
            _make_rule(
                rule_id="caption-format",
                category="图表公式",
                description="图题和表题格式",
                selector={"element_type": "paragraph", "exclude_empty": True, "text_regex": "^(图|表)\\s*\\d+([-.]\\d+)*\\s*.+"},
                expected=caption_expected,
                severity="minor",
                suggestion="按学校规范统一图题、表题格式。",
                auto_fix=False,
            )
        )

    reference_context = _context_for_best_label(spec_text, ("参考文献”内容", "参考文献内容", "参考文献", "文献列表"))
    reference_expected = _extract_paragraph_expected(reference_context)
    if reference_expected and global_line_spacing is not None:
        reference_expected.setdefault("line_spacing", global_line_spacing)
    if "悬挂" in reference_context and "first_line_indent_cm" not in reference_expected:
        reference_expected["first_line_indent_cm"] = -0.74
    if reference_expected:
        rules.append(
            _make_rule(
                rule_id="references-format",
                category="参考文献",
                description="参考文献条目格式",
                selector={
                    "element_type": "paragraph",
                    "exclude_empty": True,
                    "text_regex": "^(\\s*(\\[\\d{1,3}\\]|\\d{1,3}[.)、])\\s*.+|\\s*\\d{1,3}\\s+.+)",
                },
                expected=reference_expected,
                severity="major",
                suggestion="按学校规范统一参考文献条目格式。",
                auto_fix=False,
            )
        )

    required_sections = _extract_required_sections(spec_text)
    manual_review_notes = _extract_manual_review_notes(spec_text)
    if manual_review_notes:
        notes.append("规范中包含需人工或渲染校验确认的项目，已写入 manual_review_notes。")

    profile = {
        "profile_id": profile_id,
        "display_name": display_name,
        "version": version,
        "is_demo": False,
        "is_template": False,
        "is_draft": False,
        "source": source,
        "description": "根据上传的学校官方规范自动生成；LLM 不可用或未返回 JSON 时由本地确定性抽取器生成。",
        "rules": rules,
        "required_sections": required_sections,
        "manual_review_notes": manual_review_notes,
        "normalization": {
            "method": "local_deterministic_extractor",
            "supported_expected_fields": sorted(SUPPORTED_EXPECTED_FIELDS),
            "notes": notes,
        },
    }
    return profile, notes


def _make_rule(
    *,
    rule_id: str,
    category: str,
    description: str,
    selector: dict[str, Any],
    expected: dict[str, Any],
    severity: str,
    suggestion: str,
    auto_fix: bool,
    safe_paragraph_auto_fix: bool = False,
    safe_fix_fields: list[str] | None = None,
) -> dict[str, Any]:
    rule = {
        "id": rule_id,
        "category": category,
        "description": description,
        "selector": selector,
        "expected": expected,
        "severity": severity,
        "status": "auto_fixable" if auto_fix else "manual_guided",
        "confidence": 0.94,
        "auto_fix": auto_fix,
        "tolerance": 0.08,
        "suggestion": suggestion,
    }
    if safe_paragraph_auto_fix:
        rule["safe_paragraph_auto_fix"] = True
        rule["safe_fix_fields"] = safe_fix_fields or []
    return rule


def _is_body_rule(rule: dict[str, Any]) -> bool:
    text = " ".join(str(rule.get(key, "")) for key in ("id", "category", "description"))
    return "body" in text.lower() or "正文" in text


def _safe_paragraph_fix_fields(raw_fields: Any) -> list[str]:
    allowed = ["line_spacing", "first_line_indent_cm", "space_before_pt", "space_after_pt"]
    if not isinstance(raw_fields, list) or not raw_fields:
        return allowed
    selected = [field for field in raw_fields if field in allowed]
    return selected or allowed


def _with_default_body_exclusions(selector: dict[str, Any]) -> dict[str, Any]:
    result = dict(selector)
    patterns = list(result.get("text_regex_not") or [])
    for pattern in [
        "^(第[一二三四五六七八九十0-9]+章|[一二三四五六七八九十]+、)\\s*.+",
        "^(\\d+\\.\\d+|（[一二三四五六七八九十]+）)\\s*.+",
        "^摘要$",
        "^Abstract$",
        "^关键词[:：]",
        "^Key\\s*words?\\s*[:：]",
        "^Keywords\\s*[:：]",
        "^目录$",
        "^Contents$",
        "^参考文献$",
        "^致谢$",
        "^(图|表)\\s*\\d+",
    ]:
        if pattern not in patterns:
            patterns.append(pattern)
    result["text_regex_not"] = patterns
    return result


def _extract_section_expected(text: str) -> dict[str, Any]:
    expected: dict[str, Any] = {}
    if re.search(r"\bA4\b|A４|210\s*mm|29\.7\s*cm|297\s*mm", text, flags=re.IGNORECASE):
        expected["page_width_cm"] = 21.0
        expected["page_height_cm"] = 29.7
    expected.update(_extract_margin_values(text))
    return expected


def _extract_margin_values(text: str) -> dict[str, float]:
    expected: dict[str, float] = {}
    side_map = {
        "上": "top_margin_cm",
        "顶": "top_margin_cm",
        "下": "bottom_margin_cm",
        "底": "bottom_margin_cm",
        "左": "left_margin_cm",
        "右": "right_margin_cm",
    }
    compact = text.replace("．", ".")
    for match in re.finditer(r"(上|下|左|右|顶端|底端|顶部|底部)(?:边距|页边距)?[为:：]?\s*(\d+(?:\.\d+)?)\s*(厘米|cm|mm|毫米)", compact, flags=re.IGNORECASE):
        side, value, unit = match.groups()
        key = side_map[side[0]]
        expected[key] = _unit_to_cm(float(value), unit)

    pair_patterns = [
        (r"(?:上[、和及]?下|上下)(?:边距|页边距)?[为:：]?\s*(\d+(?:\.\d+)?)\s*(厘米|cm|mm|毫米)", ("top_margin_cm", "bottom_margin_cm")),
        (r"(?:左[、和及]?右|左右)(?:边距|页边距)?[为:：]?\s*(\d+(?:\.\d+)?)\s*(厘米|cm|mm|毫米)", ("left_margin_cm", "right_margin_cm")),
    ]
    for pattern, keys in pair_patterns:
        match = re.search(pattern, compact, flags=re.IGNORECASE)
        if not match:
            continue
        value, unit = match.groups()
        cm = _unit_to_cm(float(value), unit)
        for key in keys:
            expected.setdefault(key, cm)
    return expected


def _extract_paragraph_expected(context: str) -> dict[str, Any]:
    expected: dict[str, Any] = {}
    if not context:
        return expected
    font_name = _extract_font_name(context)
    if font_name:
        expected["font_name"] = font_name
    font_size = _extract_font_size(context)
    if font_size is not None:
        expected["font_size_pt"] = font_size
    line_spacing = _extract_line_spacing(context)
    if line_spacing is not None:
        expected["line_spacing"] = line_spacing
    indent = _extract_first_line_indent(context)
    if indent is not None:
        expected["first_line_indent_cm"] = indent
    space_before = _extract_space_pt(context, "段前")
    if space_before is not None:
        expected["space_before_pt"] = space_before
    space_after = _extract_space_pt(context, "段后")
    if space_after is not None:
        expected["space_after_pt"] = space_after
    alignment = _extract_alignment(context)
    if alignment:
        expected["alignment"] = alignment
    if "加粗" in context or "粗体" in context or "黑体" in context:
        expected["bold"] = True
    if "斜体" in context:
        expected["italic"] = True
    return expected


def _extract_font_name(text: str) -> str | None:
    for font in CHINESE_FONTS:
        if font.lower() in text.lower():
            return font
    return None


def _extract_font_size(text: str) -> float | None:
    for name in sorted(FONT_SIZE_NAME_TO_PT, key=len, reverse=True):
        if re.search(fr"{re.escape(name)}\s*号", text):
            return FONT_SIZE_NAME_TO_PT[name]
    match = re.search(r"(\d+(?:\.\d+)?)\s*(?:pt|磅)", text, flags=re.IGNORECASE)
    if match:
        return float(match.group(1))
    return None


def _extract_line_spacing(text: str) -> float | None:
    match = re.search(r"(\d+(?:\.\d+)?)\s*倍(?:行距)?", text)
    if match:
        return float(match.group(1))
    if "单倍行距" in text:
        return 1.0
    if "两倍行距" in text or "2倍行距" in text:
        return 2.0
    return None


def _extract_global_line_spacing(text: str) -> float | None:
    for line in text.splitlines():
        if "未标注行距" in line or "未标明行距" in line or "以上未标注" in line:
            value = _extract_line_spacing(line)
            if value is not None:
                return value
    return None


def _extract_first_line_indent(text: str) -> float | None:
    if "无首行缩进" in text or "首行不缩进" in text:
        return 0
    match = re.search(r"首行缩进\s*(\d+(?:\.\d+)?)\s*(厘米|cm|mm|毫米)", text, flags=re.IGNORECASE)
    if match:
        value, unit = match.groups()
        return _unit_to_cm(float(value), unit)
    if re.search(r"首行缩进\s*2\s*(?:个)?汉?字符|首行缩进\s*两\s*(?:个)?汉?字符|首行缩进\s*2\s*字", text):
        return 0.74
    return None


def _extract_space_pt(text: str, label: str) -> float | None:
    match = re.search(fr"{label}\s*(\d+(?:\.\d+)?)\s*(?:pt|磅)", text, flags=re.IGNORECASE)
    if match:
        return float(match.group(1))
    if re.search(fr"{label}\s*0", text):
        return 0
    return None


def _extract_alignment(text: str) -> str | None:
    if "居中" in text:
        return "center"
    if "两端对齐" in text or "端对齐" in text:
        return "justify"
    if "右对齐" in text:
        return "right"
    if "左对齐" in text:
        return "left"
    return None


def _extract_required_sections(text: str) -> list[dict[str, Any]]:
    specs = [
        ("required-abstract", "特殊模块", "摘要", "^摘要$", "major", "论文应包含摘要模块。"),
        ("required-keywords", "特殊模块", "关键词", "^关键词[:：]", "major", "摘要后应包含关键词。"),
        ("required-toc", "目录", "目录", "^目录$", "major", "论文应包含目录。"),
        ("required-references", "参考文献", "参考文献", "^参考文献$", "critical", "论文应包含参考文献模块。"),
        ("required-thanks", "特殊模块", "致谢", "^致谢$", "minor", "如学校规范要求，应包含致谢模块。"),
        ("required-appendix", "特殊模块", "附录", "^附录", "minor", "如学校规范要求，应包含附录模块。"),
    ]
    sections: list[dict[str, Any]] = []
    for section_id, category, label, pattern, severity, suggestion in specs:
        if label not in text:
            continue
        sections.append(
            {
                "id": section_id,
                "category": category,
                "label": label,
                "text_regex": pattern,
                "severity": severity,
                "suggestion": suggestion,
            }
        )
    return sections


def _extract_manual_review_notes(text: str) -> list[str]:
    notes: list[str] = []
    checks = [
        ("目录" in text and "页码" in text, "目录页码、页码连续性需要渲染校验或人工复核。"),
        ("页眉" in text or "页脚" in text, "页眉页脚涉及节、奇偶页或首页差异时需要人工复核。"),
        ("跨页" in text, "跨页表格或跨页版式需要人工复核。"),
        ("图片" in text and ("清晰" in text or "分辨率" in text), "图片清晰度、压缩失真需要人工复核。"),
        ("公式" in text, "复杂公式编号和对齐需要人工复核。"),
    ]
    for enabled, note in checks:
        if enabled and note not in notes:
            notes.append(note)
    return notes


def _context_for_keywords(text: str, keywords: tuple[str, ...]) -> str:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    matched: list[str] = []
    for index, line in enumerate(lines):
        if any(keyword in line for keyword in keywords):
            start = max(0, index - 1)
            end = min(len(lines), index + 2)
            matched.extend(lines[start:end])
    return "\n".join(dict.fromkeys(matched))


def _context_for_best_label(text: str, labels: tuple[str, ...]) -> str:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    table_matches: list[str] = []
    for line in lines:
        if "|" not in line:
            continue
        left = line.split("|", 1)[0].strip().strip("“”\"")
        if any(label.strip("“”\"") == left or label in left for label in labels):
            table_matches.append(line)
    if table_matches:
        return "\n".join(table_matches)

    format_pattern = re.compile(
        r"字体|字号|号|宋体|黑体|楷体|仿宋|Times|行距|缩进|居中|对齐|段前|段后|加粗|小四|五号|三号|四号"
    )
    direct_matches = [
        line
        for line in lines
        if any(label in line for label in labels) and format_pattern.search(line)
    ]
    if direct_matches:
        return "\n".join(direct_matches)
    return _context_for_keywords(text, labels)


def _unit_to_cm(value: float, unit: str) -> float:
    if unit.lower() == "mm" or unit == "毫米":
        value = value / 10
    return round(value, 2)


def _extract_docx_text(payload: bytes) -> str:
    document = Document(io.BytesIO(payload))
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = " ".join(paragraph.text.split())
        if text:
            lines.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            row_text = " | ".join(cell for cell in cells if cell)
            if row_text:
                lines.append(row_text)
    for index, section in enumerate(document.sections, start=1):
        for label, part in (
            ("页眉", section.header),
            ("页脚", section.footer),
            ("首页页眉", section.first_page_header),
            ("首页页脚", section.first_page_footer),
            ("偶数页页眉", section.even_page_header),
            ("偶数页页脚", section.even_page_footer),
        ):
            for paragraph in part.paragraphs:
                text = " ".join(paragraph.text.split())
                if text:
                    lines.append(f"第{index}节{label}: {text}")
    return "\n".join(lines)


def _decode_text(payload: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return payload.decode(encoding)
        except UnicodeDecodeError:
            continue
    return payload.decode("utf-8", errors="replace")


def _extract_pdf_text(payload: bytes) -> str:
    try:
        from pypdf import PdfReader  # type: ignore[import-not-found]
    except ModuleNotFoundError as exc:
        raise SpecNormalizationError("PDF 自动抽取需要安装 pypdf；建议上传学校规范 DOCX/TXT/MD 版本。") from exc
    reader = PdfReader(io.BytesIO(payload))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _compact_spec_text(text: str) -> str:
    lines = [" ".join(line.split()) for line in text.splitlines()]
    lines = [line for line in lines if line]
    if not lines:
        return ""
    joined = "\n".join(lines)
    if len(joined) <= MAX_SPEC_CHARS:
        return joined

    keyword_pattern = re.compile(
        r"字体|字号|行距|页边距|装订|纸张|A4|标题|正文|摘要|关键词|目录|参考文献|致谢|附录|"
        r"页眉|页脚|页码|缩进|段前|段后|居中|左对齐|右对齐|两端对齐|图|表|公式|"
        r"宋体|黑体|楷体|仿宋|Times|厘米|cm|磅|pt|倍",
        flags=re.IGNORECASE,
    )
    selected: list[str] = []
    for line in lines:
        if keyword_pattern.search(line):
            selected.append(line)
    compacted = "\n".join(selected) or joined
    if len(compacted) <= MAX_SPEC_CHARS:
        return compacted
    return compacted[:MAX_SPEC_CHARS] + "\n[规范文本过长，已截断；请优先上传更聚焦的学校格式规范文件。]"


def _find_balanced_json_object(text: str) -> str:
    start = text.find("{")
    if start < 0:
        return ""
    depth = 0
    in_string = False
    escaped = False
    for index in range(start, len(text)):
        char = text[index]
        if in_string:
            if escaped:
                escaped = False
            elif char == "\\":
                escaped = True
            elif char == '"':
                in_string = False
            continue
        if char == '"':
            in_string = True
        elif char == "{":
            depth += 1
        elif char == "}":
            depth -= 1
            if depth == 0:
                return text[start : index + 1]
    return ""


def _selector_regex_valid(selector: dict[str, Any]) -> bool:
    pattern = selector.get("text_regex")
    if pattern and not _regex_valid(str(pattern)):
        return False
    for item in selector.get("text_regex_not") or []:
        if not _regex_valid(str(item)):
            return False
    return True


def _regex_valid(pattern: str) -> bool:
    try:
        re.compile(pattern)
    except re.error:
        return False
    return True


def _safe_rule_id(value: str, default: str) -> str:
    cleaned = re.sub(r"[^a-zA-Z0-9_\-]+", "-", value.strip()).strip("-").lower()
    return cleaned or default


def _clamp_float(value: Any, default: float, minimum: float, maximum: float) -> float:
    try:
        numeric = float(value)
    except (TypeError, ValueError):
        numeric = default
    return min(max(numeric, minimum), maximum)
