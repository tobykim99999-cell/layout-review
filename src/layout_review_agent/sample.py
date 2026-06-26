from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


def create_bad_sample(output_dir: str | Path) -> Path:
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    path = output_dir / "bad_thesis.docx"

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    document.add_paragraph("摘要")
    document.add_paragraph("关键词：排版；审核；智能体")

    title = document.add_paragraph("第一章 绪论")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.line_spacing = 1.0
    title_run = title.runs[0]
    title_run.font.name = "宋体"
    title_run.font.size = Pt(10)
    title_run.font.bold = False

    body = document.add_paragraph("这是一个用于测试排版审核智能体的正文段落，故意设置了错误的字体、字号、缩进和行距。")
    body.alignment = WD_ALIGN_PARAGRAPH.LEFT
    body.paragraph_format.first_line_indent = Cm(0)
    body.paragraph_format.line_spacing = 1.0
    body.paragraph_format.space_after = Pt(6)
    body_run = body.runs[0]
    body_run.font.name = "Calibri"
    body_run.font.size = Pt(10)

    caption = document.add_paragraph("图 1 测试图题")
    caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
    caption.runs[0].font.name = "Calibri"
    caption.runs[0].font.size = Pt(12)

    document.save(path)
    return path
